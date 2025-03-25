from flask import Flask, render_template, request, jsonify, send_file
import json
import os
import time
import datetime
import traceback
import requests  # Use requests directly instead of ollama client
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading
import subprocess
import tempfile
from pathlib import Path

# Ollama API configuration
OLLAMA_API_BASE = "http://localhost:11434"
OLLAMA_LIST_ENDPOINT = f"{OLLAMA_API_BASE}/api/tags"
OLLAMA_GENERATE_ENDPOINT = f"{OLLAMA_API_BASE}/api/generate"

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 1 * 1024 * 1024  # 1 MB max upload size

# Configuration
MODEL_NAME = "llama3.1:8b"  # Exact model name as shown in ollama list
TEMPERATURE = 0.7
MAX_TOKENS = 3000  # Increased to generate much longer content

# Ensure directories exist
os.makedirs("generated_sops", exist_ok=True)
os.makedirs("saved_data", exist_ok=True)
os.makedirs('data', exist_ok=True)
os.makedirs('templates', exist_ok=True)
os.makedirs('static', exist_ok=True)

# Define section prompts with specific content requirements and exact word limits
SOP_SECTION_PROMPTS = {
    "introduction": {
        "title": "Respected Sir/Ma'am",
        "word_limit": 54,
        "content": "Write a formal and respectful introduction paragraph for a Statement of Purpose. Address it to the admissions committee or visa officer. Mention the purpose of the letter (to apply for admission to the mentioned course at the mentioned university). Use sophisticated language but keep it EXACTLY 54 words. No more, no less."
    },
    "academic_background": {
        "title": "Academic Background",
        "word_limit": 71,
        "content": "Write about the educational background, starting from 10th grade to the highest level of education completed. Include details about board/university, percentages/CGPA, and years of completion. Keep it EXACTLY 71 words. No more, no less."
    },
    "language_proficiency": {
        "title": "Language Proficiency",
        "word_limit": 180,
        "content": "Describe English language proficiency based on IELTS/PTE scores. Mention the overall score and individual section scores. Explain how these scores demonstrate the ability to succeed in an academic environment where English is the medium of instruction. Keep it EXACTLY 180 words. No more, no less."
    },
    "financial_background": {
        "title": "Financial Background",
        "word_limit": 148,
        "content": "Explain how the education and expenses will be funded. Mention the sponsors (usually parents), their occupations, annual income, savings, and financial capacity to support the education. If applicable, mention scholarships or other funding sources. Keep it EXACTLY 148 words. No more, no less."
    },
    "why_country": {
        "title": "Why I Choose this Country for my Studies",
        "word_limit": 122,
        "content": "Explain reasons for choosing the specific country for education. Discuss the quality of education, global recognition of degrees, cultural diversity, safe environment, and opportunities for international students. Make it specific to the country mentioned. Keep it EXACTLY 122 words. No more, no less."
    },
    "career_opportunities": {
        "title": "Career Opportunities in My Country After Completing the Program",
        "word_limit": 354,
        "content": "Discuss career prospects in the home country after completing the program. Mention specific industries, job roles, and growing demand for professionals in the field of study. Emphasize intention to return to home country and contribute to its development. Keep it EXACTLY 354 words. No more, no less."
    },
    "family_ties": {
        "title": "My Family Ties and Return to Home Country",
        "word_limit": 275,
        "content": "Describe family ties and reasons for returning to home country after studies. Mention family members, family business (if any), cultural attachments, and responsibilities that ensure return to the home country. Keep it EXACTLY 275 words. No more, no less."
    },
    "conclusion": {
        "title": "Conclusion",
        "word_limit": 70,
        "content": "Provide a concise conclusion summarizing the key points of the SOP. Express gratitude for considering the application, and mention enthusiasm for joining the program. End with formal closing. Keep it EXACTLY 70 words. No more, no less."
    }
}

def check_ollama_status():
    """Check if Ollama is running and if the required model is available using direct REST API"""
    try:
        # Try to list models using REST API
        response = requests.get(OLLAMA_LIST_ENDPOINT, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            
            # Extract model names
            if 'models' in data:
                models = data['models']
                model_names = [str(model.get('name', '')) for model in models]
            else:
                # Handle older Ollama API format
                model_names = [model.get('name') for model in data.get('models', [])]
                if not model_names and 'models' not in data:
                    print("Unexpected response format:", data)
                    model_names = []
                    for model in data.get('models', []):
                        print(f"Model data: {model}")
                        if 'name' in model:
                            model_names.append(model['name'])
            
            print(f"Available models from API: {model_names}")
            print(f"Looking for model: '{MODEL_NAME}'")
            
            # Check if our model is available
            exact_match = MODEL_NAME in model_names
            flexible_match = any(MODEL_NAME in model_name for model_name in model_names)
            
            print(f"Exact match: {exact_match}, Flexible match: {flexible_match}")
            
            return {
                "ollama_running": True,
                "model_available": exact_match or flexible_match,
                "models": model_names,
                "exact_match": exact_match,
                "flexible_match": flexible_match
            }
        else:
            print(f"Ollama API returned status code {response.status_code}")
            return {
                "ollama_running": False,
                "model_available": False,
                "error": f"Ollama API returned status code {response.status_code}"
            }
            
    except requests.exceptions.RequestException as e:
        print(f"Failed to connect to Ollama API: {str(e)}")
        traceback.print_exc()
        return {
            "ollama_running": False,
            "model_available": False,
            "error": f"Failed to connect to Ollama API: {str(e)}"
        }

def generate_with_ollama_api(model, prompt, temperature=0.7, max_tokens=1000):
    """Generate text using Ollama REST API directly"""
    try:
        payload = {
            "model": model,
            "prompt": prompt,
            "options": {
                "temperature": temperature,
                "num_predict": max_tokens
            },
            "stream": False
        }
        
        print(f"Sending request to Ollama API with model: {model}")
        response = requests.post(OLLAMA_GENERATE_ENDPOINT, json=payload, timeout=120)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('response', '')
        else:
            print(f"Ollama API error: {response.status_code}, {response.text}")
            return f"Error: Ollama API returned status code {response.status_code}"
            
    except requests.exceptions.RequestException as e:
        print(f"Request to Ollama API failed: {str(e)}")
        traceback.print_exc()
        return f"Error: Failed to communicate with Ollama API - {str(e)}"

def generate_section_with_ollama(section_key, user_data):
    """Generate a section of the SOP using Ollama API"""
    section_info = SOP_SECTION_PROMPTS[section_key]
    prompt = prepare_prompt(section_key, section_info, user_data)
    
    # Calculate tokens based on target word count (approximately 1.5 tokens per word)
    # Ensure we have enough tokens for the response (at least 2x the word limit)
    section_tokens = max(section_info['word_limit'] * 2, 1000)
    
    try:
        # Check if model is available first
        status = check_ollama_status()
        if not status["ollama_running"]:
            return f"Error: Ollama service is not running. Please start Ollama and try again."
        if not status["model_available"]:
            return f"Error: Model '{MODEL_NAME}' not found. Please install it using: ollama pull {MODEL_NAME}"
        
        print(f"Generating section '{section_key}' with model: {MODEL_NAME}")
        print(f"Target word count: {section_info['word_limit']}, Tokens: {section_tokens}")
        
        # Call Ollama API directly with section-specific token limit
        generated_text = generate_with_ollama_api(
            model=MODEL_NAME,
            prompt=prompt,
            temperature=TEMPERATURE,
            max_tokens=section_tokens
        )
        
        if not generated_text or "Error:" in generated_text:
            return f"Error generating {section_info['title']}: {generated_text}"
            
        # Validate the generated content
        word_count = len(generated_text.split())
        print(f"Generated {word_count} words for section: {section_key}")
        
        if word_count < 10:
            raise Exception(f"Generated content for {section_key} is too short")
            
        return generated_text.strip()
    except Exception as e:
        print(f"Error generating {section_key}: {str(e)}")
        traceback.print_exc()
        return f"Error generating {section_info['title']}: {str(e)}"

def prepare_prompt(section_key, section_info, user_data):
    """Prepare a prompt for the section with user data"""
    
    system_prompt = f"""You are an expert Statement of Purpose (SOP) writer. Write a {section_info['title']} section for a Statement of Purpose with EXACTLY {section_info['word_limit']} words. Not one word more or less.

Requirements:
{section_info['content']}

Word Limit: EXACTLY {section_info['word_limit']} words. Count your words carefully.
Style: Formal, clear, and professional. Use first-person perspective.
Format: No headings, titles, or prefixes. Just write the content of the section.

IMPORTANT: Your response must be EXACTLY {section_info['word_limit']} words. Count the words carefully.
"""

    # Add user data to make it personalized
    user_prompt = "\nUser Information:\n"
    
    relevant_fields = {
        "introduction": ["name", "university_name", "course", "country"],
        "academic_background": ["name", "state", "tenth_board", "tenth_marks", "tenth_year", 
                              "twelfth_board", "twelfth_marks", "twelfth_year",
                              "bachelors_degree", "bachelors_college", "bachelors_cgpa"],
        "language_proficiency": ["test_type", "listening", "reading", "writing", "speaking", "overall"],
        "financial_background": ["father_income", "mother_income", "father_funds", "mother_funds", "fixed_deposits"],
        "why_country": ["country", "university_name", "course"],
        "career_opportunities": ["course", "country"],
        "family_ties": ["state", "family_members"],
        "conclusion": ["name", "university_name", "course", "country"]
    }
    
    for field in relevant_fields.get(section_key, []):
        if field in user_data and user_data[field]:
            field_name = field.replace('_', ' ').title()
            user_prompt += f"- {field_name}: {user_data[field]}\n"
    
    complete_prompt = f"{system_prompt}\n{user_prompt}\nWrite the {section_info['title']} section with EXACTLY {section_info['word_limit']} words:"
    return complete_prompt

def generate_complete_sop(user_data):
    """Generate complete SOP with all sections"""
    
    start_time = time.time()
    
    # Save user data for future reference
    save_user_data(user_data)
    
    complete_sop = ""
    sections_content = {}
    
    try:
        # Generate each section sequentially
        for section_key in SOP_SECTION_PROMPTS.keys():
            section_content = generate_section_with_ollama(section_key, user_data)
            sections_content[section_key] = section_content
            
            section_title = SOP_SECTION_PROMPTS[section_key]["title"]
            complete_sop += f"{section_title}\n\n{section_content}\n\n"
    
        # Save the generated SOP
        save_generated_sop(user_data.get('name', 'unnamed'), complete_sop, sections_content)
        
        end_time = time.time()
        generation_time = round(end_time - start_time, 2)
        
        return {"success": True, "sop_content": complete_sop, "generation_time": generation_time}
    
    except Exception as e:
        print(f"Error generating SOP: {str(e)}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def save_user_data(user_data):
    """Save user data to a JSON file"""
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    user_name = user_data.get('name', 'unnamed').replace(' ', '_')
    filename = f"data/user_data_{user_name}_{timestamp}.json"
    
    with open(filename, 'w') as f:
        json.dump(user_data, f, indent=4)
    
    return filename

def save_generated_sop(name, complete_sop, sections_content):
    """Save the generated SOP to a JSON file"""
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    user_name = name.replace(' ', '_')
    filename = f"data/sop_{user_name}_{timestamp}.json"
    
    data = {
        "complete_sop": complete_sop,
        "sections": sections_content,
        "timestamp": timestamp
    }
    
    with open(filename, 'w') as f:
        json.dump(data, f, indent=4)
    
    return filename

def generate_docx(sop_content, user_name):
    """Generate a Word document from the SOP content with improved formatting"""
    doc = Document()
    
    # Set document properties
    doc.core_properties.author = user_name
    doc.core_properties.title = f"Statement of Purpose - {user_name}"
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title with formatting
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("STATEMENT OF PURPOSE")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_run.underline = True
    
    # Add a blank line after title
    doc.add_paragraph()
    
    # Split the content into sections
    parts = sop_content.split('\n\n')
    
    for i, part in enumerate(parts):
        if i % 2 == 0 and i+1 < len(parts):  # This is a header
            # Add heading with formatting
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_run = heading.add_run(part)
            heading_run.bold = True
            heading_run.font.size = Pt(13)
            
            # Add content with formatting
            content = doc.add_paragraph()
            content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content_run = content.add_run(parts[i+1])
            content_run.font.size = Pt(13)
    
    # Save to a BytesIO object
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io

def generate_pdf_from_docx(docx_io, user_name):
    """Alternative method to generate PDF"""
    try:
        # Create a temporary directory to work in
        temp_dir = tempfile.mkdtemp()
        temp_docx_path = os.path.join(temp_dir, f"SOP_{user_name}.docx")
        
        # Save the DOCX to the temporary file
        with open(temp_docx_path, 'wb') as f:
            f.write(docx_io.getvalue())
        
        # Fallback to just returning the DOCX if we can't convert
        pdf_data = None
        pdf_io = io.BytesIO()
        
        # Try different conversion methods
        conversion_successful = False
        
        # Method 1: Try direct LibreOffice conversion if available
        try:
            if os.name == 'nt':  # Windows
                # Check if LibreOffice exists in common locations
                libreoffice_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                
                soffice_path = None
                for path in libreoffice_paths:
                    if os.path.exists(path):
                        soffice_path = path
                        break
                
                if soffice_path:
                    temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                    subprocess.run([
                        soffice_path,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', temp_dir,
                        temp_docx_path
                    ], check=True, timeout=30)
                    
                    if os.path.exists(temp_pdf_path):
                        with open(temp_pdf_path, 'rb') as f:
                            pdf_data = f.read()
                        conversion_successful = True
            else:  # Linux/Mac
                temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    temp_docx_path
                ], check=True, timeout=30)
                
                if os.path.exists(temp_pdf_path):
                    with open(temp_pdf_path, 'rb') as f:
                        pdf_data = f.read()
                    conversion_successful = True
        except Exception as e:
            print(f"LibreOffice conversion failed: {str(e)}")
        
        # If all conversion methods failed, return the original DOCX
        if not conversion_successful or not pdf_data:
            print("PDF conversion failed, returning original DOCX")
            docx_io.seek(0)
            return docx_io, 'docx'
        
        # Create PDF BytesIO from the data
        pdf_io.write(pdf_data)
        pdf_io.seek(0)
        
        # Clean up temporary files
        try:
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)
        except Exception as e:
            print(f"Error cleaning up temp files: {str(e)}")
        
        return pdf_io, 'pdf'
    
    except Exception as e:
        print(f"Error in PDF conversion: {str(e)}")
        traceback.print_exc()
        
        # Return original DOCX as fallback
        docx_io.seek(0)
        return docx_io, 'docx'

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/check_ollama', methods=['GET'])
def get_ollama_status():
    status = check_ollama_status()
    return jsonify(status)

@app.route('/generate_sop', methods=['POST'])
def generate_sop():
    try:
        # Get form data and convert to dict
        user_data = {}
        for key in request.form:
            user_data[key] = request.form.get(key)
        
        # Generate the SOP
        result = generate_complete_sop(user_data)
        return jsonify(result)
    
    except Exception as e:
        print(f"Error: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)})

@app.route('/download_docx', methods=['POST'])
def download_docx():
    try:
        sop_content = request.form.get('sop_content')
        user_name = request.form.get('name', 'Unnamed')
        
        docx_io = generate_docx(sop_content, user_name)
        
        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"SOP_{user_name.replace(' ', '_')}.docx"
        )
    
    except Exception as e:
        print(f"Error generating docx: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)})

@app.route('/download_txt', methods=['POST'])
def download_txt():
    try:
        sop_content = request.form.get('sop_content')
        user_name = request.form.get('name', 'Unnamed')
        
        # Create a BytesIO object for the text content
        text_io = io.BytesIO()
        text_io.write(sop_content.encode('utf-8'))
        text_io.seek(0)
        
        return send_file(
            text_io,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f"SOP_{user_name.replace(' ', '_')}.txt"
        )
    
    except Exception as e:
        print(f"Error generating txt: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)})

@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    try:
        sop_content = request.form.get('sop_content')
        user_name = request.form.get('name', 'Unnamed')
        
        # First generate DOCX
        docx_io = generate_docx(sop_content, user_name)
        
        try:
            # Try to convert to PDF
            result_io, result_type = generate_pdf_from_docx(docx_io, user_name)
            
            if result_type == 'pdf':
                return send_file(
                    result_io,
                    mimetype='application/pdf',
                    as_attachment=True,
                    download_name=f"SOP_{user_name.replace(' ', '_')}.pdf"
                )
            else:
                # If PDF conversion failed, return DOCX instead
                return send_file(
                    result_io,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f"SOP_{user_name.replace(' ', '_')}.docx"
                ), 200, {'X-PDF-Error': 'PDF conversion failed, providing DOCX instead'}
        except Exception as pdf_error:
            print(f"PDF processing failed: {str(pdf_error)}")
            # If PDF conversion fails, return DOCX as fallback
            docx_io.seek(0)
            return send_file(
                docx_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"SOP_{user_name.replace(' ', '_')}.docx"
            ), 200, {'X-PDF-Error': 'PDF conversion failed, providing DOCX instead'}
    
    except Exception as e:
        print(f"Error in download_pdf: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)})

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000) 