from flask import Flask, render_template, request, jsonify, send_file
import json
import os
import time
from datetime import datetime
import traceback
import ollama
import docx
import io

app = Flask(__name__)

# Configuration
MODEL_NAME = "llama3.1:8b"  # Llama 3.1 8B model

# Ensure directories exist
os.makedirs("generated_sops", exist_ok=True)
os.makedirs("saved_data", exist_ok=True)

# SOP section prompts
sop_sections = {
    "introduction": "Write a statement of purpose introduction paragraph for {name} who lives in {state}, India, applying for {course} program at {university}, {country} for the {intake} intake. Keep it formal and 45-50 words.",
    
    "academic_background": "Write an academic background paragraph for a statement of purpose mentioning: 10th grade from {tenth_board} in {tenth_year} with {tenth_marks}%, 12th grade from {twelfth_board} in {twelfth_year} with {twelfth_marks}%. {bachelors_info}. Keep it formal and 80-100 words.",
    
    "language_proficiency": "Write a language proficiency paragraph for a statement of purpose mentioning {test_type} scores: Overall {overall}, Listening {listening}, Speaking {speaking}, Writing {writing}, Reading {reading}. Emphasize language abilities for academic success. Keep it formal and 100-150 words.",
    
    "program_relevance": "Write three paragraphs (total 800-900 words) for a statement of purpose explaining why {course} is relevant to the student's background in {bachelors_degree} and {work_experience}. First paragraph should connect past education to this program, second paragraph should discuss skills to be gained, and third paragraph should discuss career prospects with statistics.",
    
    "financial_background": "Write a financial background paragraph for a statement of purpose mentioning: father earns {father_income} annually, mother earns {mother_income} annually, liquid funds of {father_funds} in father's account and {mother_funds} in mother's account, {fixed_deposits}. Emphasize financial stability to complete education. Keep it formal and 100-150 words.",
    
    "country_choice": "Write two paragraphs (400 words) for a statement of purpose explaining why {country} is an ideal choice for studying. First paragraph should discuss multiculturalism and support for international students, and second paragraph should compare with other countries highlighting advantages, universities, and quality of life.",
    
    "career_opportunities": "Write two paragraphs (600-700 words) for a statement of purpose discussing career opportunities in India after completing {course}. First paragraph should discuss long-term career growth with industry statistics, and second paragraph should mention specific job roles and companies in India with salary expectations.",
    
    "family_ties": "Write three paragraphs (350-400 words) for a statement of purpose explaining motivation to return to India after studies. First paragraph should discuss family responsibilities, second paragraph should highlight cultural ties, and third paragraph should mention professional opportunities in India.",
    
    "conclusion": "Write a concluding paragraph (120 words) for a statement of purpose summarizing intentions to study {course} in {country} and commitment to return to India after completion."
}

def generate_section_with_ollama(prompt_template, user_data):
    """Generate a section of the SOP using Ollama and the Llama 3.1 model"""
    # Prepare the prompt by filling in the template
    formatted_prompt = prepare_prompt(prompt_template, user_data)
    
    try:
        # Call Ollama API
        response = ollama.generate(
            model=MODEL_NAME,
            prompt=formatted_prompt,
            stream=False
        )
        return response['response'].strip()
    except Exception as e:
        print(f"Error generating with Ollama: {str(e)}")
        traceback.print_exc()
        return f"Error generating content: {str(e)}"

def prepare_prompt(prompt_template, user_data):
    """Prepare the prompt by filling in user data"""
    # Handle special cases for certain sections
    if "academic_background" in prompt_template:
        bachelors_info = ""
        if user_data.get("bachelors_degree") and user_data.get("bachelors_college"):
            cgpa = f" with CGPA of {user_data.get('bachelors_cgpa')}" if user_data.get("bachelors_cgpa") else ""
            bachelors_info = f"Bachelor's degree in {user_data.get('bachelors_degree')} from {user_data.get('bachelors_college')}{cgpa}"
        user_data["bachelors_info"] = bachelors_info
    
    # Format the prompt with user data
    try:
        return prompt_template.format(**user_data)
    except KeyError as e:
        print(f"Missing data for prompt formatting: {e}")
        # Return a modified template that can work with missing data
        return prompt_template

def generate_complete_sop(user_data):
    """Generate a complete SOP by generating each section"""
    sections = {
        "RESPECTED SIR/MA'AM": generate_section_with_ollama(sop_sections["introduction"], user_data),
        "ACADEMIC BACKGROUND": generate_section_with_ollama(sop_sections["academic_background"], user_data),
        "LANGUAGE PROFICIENCY": generate_section_with_ollama(sop_sections["language_proficiency"], user_data),
        "PROGRAM RELEVANCE": generate_section_with_ollama(sop_sections["program_relevance"], user_data),
        "FINANCIAL BACKGROUND": generate_section_with_ollama(sop_sections["financial_background"], user_data),
        "WHY I CHOOSE THIS COUNTRY FOR MY STUDIES": generate_section_with_ollama(sop_sections["country_choice"], user_data),
        "CAREER OPPORTUNITIES IN MY COUNTRY AFTER COMPLETING THE PROGRAM": generate_section_with_ollama(sop_sections["career_opportunities"], user_data),
        "MY FAMILY TIES AND RETURN TO HOME COUNTRY": generate_section_with_ollama(sop_sections["family_ties"], user_data),
        "CONCLUSION": generate_section_with_ollama(sop_sections["conclusion"], user_data)
    }
    
    # Build the complete SOP
    sop = []
    for title, content in sections.items():
        sop.append(title)
        sop.append(content)
        sop.append("")  # Add empty line between sections
    
    return "\n\n".join(sop).strip()

def save_user_data(user_data):
    """Save user data to a JSON file"""
    try:
        # Create directory if it doesn't exist
        if not os.path.exists("saved_data"):
            os.makedirs("saved_data")
        
        # Save user data with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"saved_data/user_data_{timestamp}.json"
        
        with open(filename, 'w') as f:
            json.dump(user_data, f, indent=4)
        
        return True
    except Exception as e:
        print(f"Error saving user data: {str(e)}")
        return False

def generate_word_doc(sop_content, name="unnamed"):
    """Generate a Word document from SOP content"""
    try:
        # Create document
        doc = docx.Document()
        doc.add_heading('STATEMENT OF PURPOSE', 0)
        
        # Add content by paragraphs
        paragraphs = sop_content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if para.strip().upper() == para.strip():  # Check if paragraph is all caps (section heading)
                    doc.add_heading(para.strip(), level=1)
                else:
                    doc.add_paragraph(para.strip())
        
        # Save to memory buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        return None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_sop', methods=['POST'])
def generate_sop_endpoint():
    try:
        # Get form data
        user_data = request.form.to_dict()
        
        # Generate SOP
        sop_content = generate_complete_sop(user_data)
        
        # Save user data
        save_user_data(user_data)
        
        # Save generated SOP to file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name = user_data.get("name", "unnamed").replace(" ", "_")
        txt_filename = f"generated_sops/SOP_{name}_{timestamp}.txt"
        
        with open(txt_filename, 'w') as f:
            f.write(sop_content)
        
        return jsonify({
            "success": True,
            "sop_content": sop_content,
            "filename": txt_filename
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        })

@app.route('/download_docx', methods=['POST'])
def download_docx():
    try:
        # Get SOP content
        sop_content = request.form.get('sop_content')
        name = request.form.get('name', 'unnamed').replace(" ", "_")
        
        # Generate Word document
        file_stream = generate_word_doc(sop_content, name)
        
        if file_stream:
            # Create filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"SOP_{name}_{timestamp}.docx"
            
            # Return the Word document
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({
                "success": False,
                "error": "Failed to generate Word document"
            })
    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        })

@app.route('/download_txt', methods=['POST'])
def download_txt():
    try:
        # Get SOP content
        sop_content = request.form.get('sop_content')
        name = request.form.get('name', 'unnamed').replace(" ", "_")
        
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"SOP_{name}_{timestamp}.txt"
        
        # Create a memory file
        file_stream = io.BytesIO()
        file_stream.write(sop_content.encode('utf-8'))
        file_stream.seek(0)
        
        # Return the text file
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='text/plain'
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        })

@app.route('/check_ollama', methods=['GET'])
def check_ollama():
    try:
        # Try to list models to check if Ollama is running
        models = ollama.list()
        
        # Check if the required model is available
        model_available = any(model['name'].startswith(MODEL_NAME.split(':')[0]) for model in models.get('models', []))
        
        return jsonify({
            "ollama_running": True,
            "model_available": model_available
        })
    except Exception as e:
        return jsonify({
            "ollama_running": False,
            "error": str(e)
        })

if __name__ == '__main__':
    app.run(debug=True) 